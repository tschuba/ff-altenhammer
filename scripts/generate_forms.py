"""Erzeugt Einwilligungsformulare als DOCX für FF Altenhammer."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, os

OUT = os.path.join(os.path.dirname(__file__), '..', 'static', 'dokumente')

RED   = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x1C, 0x1C, 0x1C)
GRAY  = RGBColor(0x80, 0x80, 0x80)


def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    sec = doc.sections[0]
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bottom)
    sec.left_margin   = Cm(left)
    sec.right_margin  = Cm(right)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RED if level == 1 else BLACK
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def small(doc, text, color=GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p


def add_underline_field(doc, label, width_cm=None):
    """Label + Unterstrich-Zeile für ausfüllbare Felder."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    # Remove borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    label_cell = tbl.cell(0, 0)
    label_cell.paragraphs[0].clear()
    run = label_cell.paragraphs[0].add_run(label)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY

    line_cell = tbl.cell(1, 0)
    line_cell.paragraphs[0].clear()
    # Draw a simple underline via bottom border on this cell
    tc = line_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1C1C1C')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    line_cell.paragraphs[0].add_run(' ')
    line_cell.paragraphs[0].paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_col_fields(doc, left_label, right_label):
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(8)
    tbl.columns[1].width = Cm(7)

    def clear_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def add_bottom_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1C1C1C')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

    for row in tbl.rows:
        for cell in row.cells:
            clear_borders(cell)

    for label, col_idx in [(left_label, 0), (right_label, 1)]:
        c = tbl.cell(0, col_idx)
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(label)
        run.font.size = Pt(8.5)
        run.font.color.rgb = GRAY

        c2 = tbl.cell(1, col_idx)
        c2.paragraphs[0].clear()
        c2.paragraphs[0].add_run(' ')
        c2.paragraphs[0].paragraph_format.space_after = Pt(8)
        add_bottom_border(c2)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_checkbox_line(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 if indent else 0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run('☐  ')
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = BLACK
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────────────────────
# Formular 1: Erwachsene
# ──────────────────────────────────────────────────────────────────────────────

def build_erwachsene():
    doc = Document()
    set_margins(doc)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Freiwillige Feuerwehr Altenhammer')
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Feuerwehrverein Altenhammer e.V.  ·  VR 683  ·  Amtsgericht Weiden i.d.OPf.')
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Aufnahmeantrag und Einwilligungserklärung')
    r3.bold = True
    r3.font.size = Pt(13)
    r3.font.color.rgb = BLACK
    p3.paragraph_format.space_after = Pt(2)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('für Personen ab 18 Jahren')
    r4.font.size = Pt(10)
    r4.font.color.rgb = GRAY
    p4.paragraph_format.space_after = Pt(16)

    add_divider(doc)

    # 1. Persönliche Daten
    heading(doc, '1. Persönliche Daten')
    add_underline_field(doc, 'Vorname')
    add_underline_field(doc, 'Nachname')
    add_two_col_fields(doc, 'Geburtsdatum (TT.MM.JJJJ)', 'Staatsangehörigkeit')
    add_underline_field(doc, 'Straße und Hausnummer')
    add_two_col_fields(doc, 'PLZ', 'Ort')
    add_underline_field(doc, 'Telefon / Mobil')
    add_underline_field(doc, 'E-Mail-Adresse')

    # 2. Mitgliedschaft
    heading(doc, '2. Art der Mitgliedschaft')
    add_checkbox_line(doc, 'Aktives Mitglied  (Einsatzdienst, Übungen, Ausbildung — ab 16 Jahren)')
    add_checkbox_line(doc, 'Förderndes Mitglied  (finanzielle Unterstützung, kein Einsatzdienst)')

    # 3. Datenverarbeitung
    heading(doc, '3. Einwilligung zur Datenverarbeitung (DSGVO)')
    body(doc,
         'Ich erkläre mich einverstanden, dass meine oben angegebenen personenbezogenen Daten '
         'vom Feuerwehrverein Altenhammer e.V. zum Zweck der Mitgliederverwaltung gespeichert '
         'und verarbeitet werden. Eine Weitergabe an Dritte erfolgt nicht, ausgenommen an den '
         'Kreisfeuerwehrverband und den Bayerischen Feuerwehrverband, soweit dies für die '
         'Mitgliedschaft erforderlich ist.',
         size=9.5)
    small(doc,
          'Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung). '
          'Widerruf jederzeit möglich durch schriftliche Mitteilung an ff-altenhammer@outlook.com.')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_checkbox_line(doc, 'Ich stimme der Speicherung meiner Daten zu den genannten Zwecken zu.')

    # 4. Foto/Video
    heading(doc, '4. Einwilligung zur Veröffentlichung von Fotos und Videos (freiwillig)')
    body(doc,
         'Ich erkläre mich einverstanden, dass Fotos und Videos, auf denen ich abgebildet bin '
         'und die im Rahmen von Feuerwehr-Veranstaltungen entstehen, auf der Homepage und den '
         'Social-Media-Kanälen (Facebook, Instagram) der FF Altenhammer veröffentlicht werden dürfen.',
         size=9.5)
    small(doc, 'Diese Einwilligung ist freiwillig und kann jederzeit widerrufen werden.')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_checkbox_line(doc, 'Ich stimme der Veröffentlichung von Fotos/Videos meiner Person zu.')
    add_checkbox_line(doc, 'Ich stimme der Veröffentlichung NICHT zu.')

    # 5. Unterschrift
    heading(doc, '5. Unterschrift')
    add_two_col_fields(doc, 'Ort, Datum', 'Unterschrift')
    small(doc, 'Bitte ausgefülltes Formular beim Vorstand abgeben oder per E-Mail an ff-altenhammer@outlook.com einsenden.')

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Formular 2: Minderjährige
# ──────────────────────────────────────────────────────────────────────────────

def build_minderjaehrige():
    doc = Document()
    set_margins(doc)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Freiwillige Feuerwehr Altenhammer')
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Feuerwehrverein Altenhammer e.V.  ·  VR 683  ·  Amtsgericht Weiden i.d.OPf.')
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Aufnahmeantrag und Einwilligungserklärung')
    r3.bold = True
    r3.font.size = Pt(13)
    r3.font.color.rgb = BLACK
    p3.paragraph_format.space_after = Pt(2)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('für Minderjährige (unter 18 Jahren) — mit Einwilligung der Erziehungsberechtigten')
    r4.font.size = Pt(10)
    r4.font.color.rgb = GRAY
    p4.paragraph_format.space_after = Pt(16)

    add_divider(doc)

    # 1. Daten Kind
    heading(doc, '1. Daten des Kindes / Jugendlichen')
    add_underline_field(doc, 'Vorname')
    add_underline_field(doc, 'Nachname')
    add_two_col_fields(doc, 'Geburtsdatum (TT.MM.JJJJ)', 'Alter')
    add_underline_field(doc, 'Straße und Hausnummer')
    add_two_col_fields(doc, 'PLZ', 'Ort')
    add_underline_field(doc, 'Telefon / Mobil (des Kindes, falls vorhanden)')

    # 2. Daten Erziehungsberechtigte
    heading(doc, '2. Daten der Erziehungsberechtigten')
    add_underline_field(doc, 'Name (Erziehungsberechtigte/r)')
    add_underline_field(doc, 'Telefon / Mobil')
    add_underline_field(doc, 'E-Mail-Adresse')

    # 3. Mitgliedschaft
    heading(doc, '3. Art der Mitgliedschaft')
    add_checkbox_line(doc, 'Kinderfeuerwehr  (ab 6 Jahren, spielerisch, keine Einsätze)')
    add_checkbox_line(doc, 'Jugendfeuerwehr / Aktives Mitglied  (ab 12/16 Jahren, Übungen und Ausbildung)')
    add_checkbox_line(doc, 'Förderndes Mitglied')

    # 4. Datenverarbeitung
    heading(doc, '4. Einwilligung zur Datenverarbeitung (DSGVO)')
    body(doc,
         'Als Erziehungsberechtigte/r erkläre ich mich einverstanden, dass die personenbezogenen Daten '
         'meines Kindes vom Feuerwehrverein Altenhammer e.V. zum Zweck der Mitgliederverwaltung gespeichert '
         'und verarbeitet werden. Eine Weitergabe an Dritte erfolgt nicht, ausgenommen an den '
         'Kreisfeuerwehrverband und den Bayerischen Feuerwehrverband, soweit dies für die '
         'Mitgliedschaft erforderlich ist.',
         size=9.5)
    small(doc,
          'Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO. '
          'Widerruf jederzeit möglich durch schriftliche Mitteilung an ff-altenhammer@outlook.com.')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_checkbox_line(doc, 'Ich stimme der Speicherung der Daten meines Kindes zu den genannten Zwecken zu.')

    # 5. Foto/Video
    heading(doc, '5. Einwilligung zur Veröffentlichung von Fotos und Videos (freiwillig)')
    body(doc,
         'Ich erkläre mich einverstanden, dass Fotos und Videos, auf denen mein Kind abgebildet ist '
         'und die im Rahmen von Feuerwehr-Veranstaltungen entstehen, auf der Homepage und den '
         'Social-Media-Kanälen (Facebook, Instagram) der FF Altenhammer veröffentlicht werden dürfen.',
         size=9.5)
    small(doc, 'Diese Einwilligung ist freiwillig und kann jederzeit widerrufen werden.')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_checkbox_line(doc, 'Ich stimme der Veröffentlichung von Fotos/Videos meines Kindes zu.')
    add_checkbox_line(doc, 'Ich stimme der Veröffentlichung NICHT zu.')

    # 6. Unterschrift
    heading(doc, '6. Unterschrift der Erziehungsberechtigten')
    add_two_col_fields(doc, 'Ort, Datum', 'Unterschrift Erziehungsberechtigte/r')
    small(doc, 'Bitte ausgefülltes Formular beim Vorstand abgeben oder per E-Mail an ff-altenhammer@outlook.com einsenden.')

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Hauptprogramm
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    os.makedirs(OUT, exist_ok=True)

    paths = {
        'aufnahmeantrag-erwachsene.docx': build_erwachsene(),
        'aufnahmeantrag-minderjaehrige.docx': build_minderjaehrige(),
    }

    for filename, doc in paths.items():
        docx_path = os.path.join(OUT, filename)
        doc.save(docx_path)
        print(f'✓ {filename}')

    # PDF via LibreOffice
    print('Konvertiere zu PDF ...')
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', OUT] +
        [os.path.join(OUT, f) for f in paths],
        capture_output=True, text=True
    )
    if result.returncode == 0:
        for filename in paths:
            pdf = filename.replace('.docx', '.pdf')
            print(f'✓ {pdf}')
    else:
        print('LibreOffice Fehler:', result.stderr)
